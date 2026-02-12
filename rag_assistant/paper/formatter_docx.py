"""
Format paper drafts as Word documents (.docx).

Supports custom templates or generates clean .docx with IEEE-like styles.
"""

from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .engine import DraftedSection
from .outline import PaperOutline


class DocXFormatter:
    """Generate Word documents for drafted papers."""
    
    @staticmethod
    def format_paper_with_template(outline: PaperOutline,
                                   sections: List[DraftedSection],
                                   citations: List[Dict[str, Any]],
                                   template_path: str) -> Document:
        """
        Format paper using provided template.
        
        Args:
            outline: PaperOutline
            sections: List of DraftedSection
            citations: List of citations
            template_path: Path to .docx template
            
        Returns:
            Document object
        """
        doc = Document(template_path)
        
        # Add content to template
        DocXFormatter._add_content_to_doc(doc, outline, sections, citations)
        
        return doc
    
    @staticmethod
    def format_paper_clean(outline: PaperOutline,
                          sections: List[DraftedSection],
                          citations: List[Dict[str, Any]]) -> Document:
        """
        Generate clean .docx with IEEE-like styles (fallback).
        
        Args:
            outline: PaperOutline
            sections: List of DraftedSection
            citations: List of citations
            
        Returns:
            Document object
        """
        doc = Document()
        
        # Set up styles
        DocXFormatter._setup_styles(doc)
        
        # Add title
        title_para = doc.add_paragraph(outline.title, style='Title')
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add author/affiliation placeholder
        author_para = doc.add_paragraph("Author Name\nYour Institution")
        author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        DocXFormatter._set_font_size(author_para, 11)
        
        # Add abstract
        if outline.abstract:
            doc.add_heading('Abstract', level=1)
            doc.add_paragraph(outline.abstract)
        
        # Add sections
        DocXFormatter._add_content_to_doc(doc, outline, sections, citations)
        
        return doc
    
    @staticmethod
    def _add_content_to_doc(doc: Document, outline: PaperOutline,
                           sections: List[DraftedSection],
                           citations: List[Dict[str, Any]]):
        """Add drafted sections to document."""
        for section in sections:
            # Section heading
            heading_style = f'Heading {section.level}'
            doc.add_heading(section.title, level=section.level)
            
            # Section content
            content_para = doc.add_paragraph(section.content)
            DocXFormatter._set_font_size(content_para, 11)
            
            # Citation info
            if section.citations:
                citation_text = f"Sources: {', '.join([c['source_path'] for c in section.citations[:3]])}"
                cite_para = doc.add_paragraph(citation_text, style='List Bullet')
                DocXFormatter._set_font_size(cite_para, 9)
                DocXFormatter._set_font_color(cite_para, (128, 128, 128))
        
        # Add References section
        doc.add_heading('References', level=1)
        DocXFormatter._add_references_section(doc, citations)
    
    @staticmethod
    def _add_references_section(doc: Document, citations: List[Dict[str, Any]]):
        """Add formatted references section."""
        for i, citation in enumerate(citations, 1):
            ref_text = f"[{i}] {citation.get('source_path', 'Unknown')} - {citation.get('page_or_section', 'Section')}"
            doc.add_paragraph(ref_text, style='List Number')
    
    @staticmethod
    def _setup_styles(doc: Document):
        """Set up IEEE-like styles in document."""
        styles = doc.styles
        
        # Title style
        try:
            title_style = styles['Title']
            title_style.font.size = Pt(18)
            title_style.font.bold = True
        except:
            pass
        
        # Heading styles
        for level in range(1, 4):
            try:
                heading_style = styles[f'Heading {level}']
                heading_style.font.size = Pt(14 - (level - 1) * 2)
                heading_style.font.bold = True
            except:
                pass
    
    @staticmethod
    def _set_font_size(paragraph, size: int):
        """Set font size for all runs in paragraph."""
        for run in paragraph.runs:
            run.font.size = Pt(size)
    
    @staticmethod
    def _set_font_color(paragraph, color: Tuple[int, int, int]):
        """Set font color for all runs in paragraph."""
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(*color)
    
    @staticmethod
    def save_to_file(doc: Document, output_path: str):
        """
        Save document to file.
        
        Args:
            doc: Document object
            output_path: Path to save .docx file
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))